"""CV Tailor — reads CV/cover/job, tailors via the configured LLM, writes PDF/DOCX.

Uses the simple chat_completion() helper (not Strands tool-calling) so it works
with any OpenAI-compatible model — including ones that don't support function calls.
"""

import asyncio
import logging
import os
import re
import textwrap
from concurrent.futures import ThreadPoolExecutor
from typing import Callable, Optional

from hub.core.llm import chat_completion
from hub.core.prompts import load_prompt

logger = logging.getLogger(__name__)


class CVTailorAgent:
    """Tailors CV and cover letter to a job description using the configured LLM."""

    def __init__(self, settings=None):
        # `settings` kept for call-site compatibility; the LLM is configured via hub.config
        self.system = load_prompt("cv_tailor")

    # ── reading source files (no Strands tools) ───────────────────────

    @staticmethod
    def _read_text_file(path: str) -> str:
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                with open(path, encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _read_pdf(path: str) -> str:
        import pymupdf

        doc = pymupdf.open(path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages)

    @staticmethod
    def _read_source(path: str) -> str:
        if str(path).lower().endswith(".pdf"):
            return CVTailorAgent._read_pdf(path)
        return CVTailorAgent._read_text_file(path)

    @staticmethod
    def _clean(md: str) -> str:
        """Strip a surrounding ```markdown ... ``` fence if the model added one."""
        s = (md or "").strip()
        if s.startswith("```"):
            lines = s.splitlines()
            if lines and lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            s = "\n".join(lines).strip()
        return s

    async def _generate(self, cv_text: str, cover_text: str, job_text: str) -> dict:
        src = (
            f"=== CURRENT CV ===\n{cv_text}\n\n"
            f"=== CURRENT COVER LETTER ===\n{cover_text}\n\n"
            f"=== JOB DESCRIPTION ===\n{job_text}\n"
        )
        specs = {
            "en_cv": "Produce the tailored ENGLISH version of the candidate's CV as clean markdown. Follow the CV Formatting rules strictly. Output ONLY the CV markdown — no preamble, no code fences.",
            "en_cover": "Produce the tailored ENGLISH cover letter for this role as clean markdown. Output ONLY the letter — no preamble.",
            "de_cv": "Produce the tailored GERMAN (Deutsch) CV as clean markdown, professionally adapted for the German job market. Keep all factual details accurate. Output ONLY the German CV markdown.",
            "de_cover": "Produce the tailored GERMAN (Deutsch) cover letter as clean markdown. Output ONLY the German letter.",
        }

        async def one(key: str, task: str):
            return key, self._clean(await chat_completion(self.system, src + "\n" + task))

        return dict(await asyncio.gather(*(one(k, t) for k, t in specs.items())))

    @staticmethod
    def _next_version(output_dir: str) -> int:
        if not os.path.isdir(output_dir):
            return 1
        pattern = re.compile(r"tailored_cv_v(\d+)")
        max_ver = 0
        for name in os.listdir(output_dir):
            m = pattern.search(name)
            if m:
                max_ver = max(max_ver, int(m.group(1)))
        return max_ver + 1

    def tailor(
        self,
        cv_path: str,
        cover_letter_path: str,
        job_desc_path: str,
        output_dir: str = "output",
        photo_path: Optional[str] = None,
        progress_fn: Optional[Callable[[str], None]] = None,
    ) -> str:
        def _notify(msg: str) -> None:
            if progress_fn:
                progress_fn(msg)

        version = self._next_version(output_dir)
        cv_out = f"{output_dir}/tailored_cv_v{version}.md"
        cl_out = f"{output_dir}/tailored_cover_letter_v{version}.md"
        cv_de_out = f"{output_dir}/tailored_cv_de_v{version}.md"
        cl_de_out = f"{output_dir}/tailored_cover_letter_de_v{version}.md"

        _notify("Reading CV, cover letter and job description...")
        cv_text = self._read_source(cv_path)
        cover_text = self._read_source(cover_letter_path)
        job_text = self._read_text_file(job_desc_path)

        _notify("Generating tailored CV and cover letter (EN + DE)...")
        out = asyncio.run(self._generate(cv_text, cover_text, job_text))

        for path, key in (
            (cv_out, "en_cv"),
            (cl_out, "en_cover"),
            (cv_de_out, "de_cv"),
            (cl_de_out, "de_cover"),
        ):
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write(out.get(key, "").strip() + "\n")

        # Convert markdown outputs to PDF in parallel
        _notify("Converting to PDF...")
        conversions = [
            (self._md_to_pdf, cv_out, cv_out.replace(".md", ".pdf"), photo_path),
            (self._md_to_pdf, cl_out, cl_out.replace(".md", ".pdf"), None),
            (self._md_to_pdf, cv_de_out, cv_de_out.replace(".md", ".pdf"), photo_path),
            (self._md_to_pdf, cl_de_out, cl_de_out.replace(".md", ".pdf"), None),
        ]

        with ThreadPoolExecutor(max_workers=4) as pool:
            futures = [
                pool.submit(fn, src, dst, photo_path=photo)
                for fn, src, dst, photo in conversions
            ]
            for f in futures:
                f.result()  # propagate exceptions

        _notify("Files ready!")
        return (
            "Tailored your CV and cover letter to the job description (English + German) "
            "and saved them as PDF under output/."
        )

    # ── MD → DOCX conversion ──────────────────────────────────────────

    @staticmethod
    def _normalize_date_line(line: str) -> str:
        months = (
            "January|February|March|April|May|June|July|"
            "August|September|October|November|December"
        )
        return re.sub(rf"(\d{{4}})({months})", r"\1 – \2", line)

    @staticmethod
    def _strip_link_labels(line: str) -> str:
        line = re.sub(r"\s*\[LinkedIn\]\s*", " ", line, flags=re.IGNORECASE)
        line = re.sub(r"\s*\[github\]\s*", " ", line, flags=re.IGNORECASE)
        return re.sub(r"\s+", " ", line).strip()

    @staticmethod
    def _md_to_docx(md_path: str, docx_path: str, photo_path: Optional[str] = None) -> None:
        if not os.path.exists(md_path):
            return
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
                try:
                    with open(md_path, encoding=enc) as f:
                        raw_lines = f.readlines()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(md_path, encoding="utf-8", errors="replace") as f:
                    raw_lines = f.readlines()
            lines = [
                CVTailorAgent._strip_link_labels(CVTailorAgent._normalize_date_line(ln))
                for ln in raw_lines
            ]

            doc = Document()
            style = doc.styles["Normal"]
            style.paragraph_format.space_after = Pt(6)

            if photo_path and os.path.exists(photo_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(photo_path, width=Pt(90), height=Pt(110))
                except Exception:
                    pass

            for line in lines:
                clean = line.rstrip("\n")
                if clean.startswith("# "):
                    doc.add_heading(clean[2:].strip(), level=0)
                elif clean.startswith("## "):
                    doc.add_heading(clean[3:].strip().upper().replace("**", ""), level=1)
                elif clean.startswith("### "):
                    rest = clean[4:].replace("**", "").strip()
                    if "\t" in rest:
                        title_part, location_part = rest.split("\t", 1)
                        p = doc.add_paragraph()
                        run = p.add_run(title_part.strip() + "  ")
                        run.bold = True
                        run = p.add_run(location_part.strip())
                        run.bold = True
                        p.paragraph_format.space_before = Pt(10)
                    else:
                        p = doc.add_paragraph(rest)
                        p.paragraph_format.space_before = Pt(10)
                        for run in p.runs:
                            run.bold = True
                elif clean.startswith("---"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(8)
                elif clean.startswith("- ") or clean.startswith("* "):
                    text = clean[2:]
                    _add_paragraph_with_bold(doc, text, style="List Bullet")
                else:
                    _add_paragraph_with_bold(doc, clean)

            os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
            doc.save(docx_path)
        except Exception as e:
            logger.exception("Word generation failed for %s: %s", md_path, e)
            raise

    @staticmethod
    def _md_to_pdf(md_path: str, pdf_path: str, photo_path: Optional[str] = None) -> None:
        if not os.path.exists(md_path):
            return
        try:
            import pymupdf

            for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
                try:
                    with open(md_path, encoding=enc) as f:
                        raw_lines = f.readlines()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(md_path, encoding="utf-8", errors="replace") as f:
                    raw_lines = f.readlines()
            lines = [
                CVTailorAgent._strip_link_labels(CVTailorAgent._normalize_date_line(ln))
                for ln in raw_lines
            ]

            doc = pymupdf.open()
            fontsize = 9.5
            line_height = fontsize * 1.4
            margin = 50
            page_w, page_h = 595, 842
            bottom_y = page_h - margin
            img_w, img_h = 90, 110
            photo_top = margin - 8
            photo_rect = pymupdf.Rect(page_w - margin - img_w, photo_top, page_w - margin, photo_top + img_h)
            first_sep_min_y = photo_top + img_h + 12
            full_width = page_w - 2 * margin

            def chars_fit(w, sz):
                return max(20, int(w / (0.55 * sz)))

            page = None
            y = margin

            for line in lines:
                clean = line.rstrip("\n")
                if page is None:
                    page = doc.new_page(width=page_w, height=page_h)
                    y = margin
                    if photo_path and os.path.exists(photo_path):
                        page.insert_image(photo_rect, filename=photo_path, keep_proportion=True, overlay=False)

                has_photo = photo_path and os.path.exists(photo_path)
                cw = (full_width - img_w - 15) if has_photo and page.number == 0 else full_width
                right = margin + cw

                def new_page_if_needed():
                    nonlocal page, y
                    if y >= bottom_y:
                        page = doc.new_page(width=page_w, height=page_h)
                        y = margin

                if clean.startswith("# "):
                    for part in textwrap.wrap(clean[2:], width=chars_fit(cw, 14)):
                        new_page_if_needed()
                        page.insert_text((margin, y), part, fontsize=14, fontname="helv")
                        y += line_height
                elif clean.startswith("## "):
                    new_page_if_needed()
                    text = clean[3:].strip().upper().replace("**", "")
                    for part in textwrap.wrap(text, width=chars_fit(cw, 13)):
                        new_page_if_needed()
                        page.insert_text((margin, y), part, fontsize=13, fontname="hebo")
                        y += line_height
                    y += 8
                elif clean.startswith("### "):
                    new_page_if_needed()
                    rest = clean[4:]
                    if "\t" in rest:
                        title = rest.split("\t", 1)[0].replace("**", "").strip()
                        loc = rest.split("\t", 1)[1].replace("**", "").strip()
                        page.insert_text((margin, y), title, fontsize=10.5, fontname="hebo")
                        loc_w = pymupdf.get_text_length(loc, fontname="hebo", fontsize=10.5)
                        page.insert_text((right - loc_w, y), loc, fontsize=10.5, fontname="hebo")
                        y += line_height + 6
                    else:
                        text = rest.replace("**", "").strip()
                        for part in textwrap.wrap(text, width=chars_fit(cw, 10.5)):
                            new_page_if_needed()
                            page.insert_text((margin, y), part, fontsize=10.5, fontname="hebo")
                            y += line_height + 6
                elif clean.startswith("---"):
                    new_page_if_needed()
                    if page.number == 0 and has_photo:
                        y = max(y, first_sep_min_y)
                    page.draw_line((margin, y - 2), (right, y - 2))
                    y += line_height + 6
                elif clean.startswith(("- ", "* ")):
                    text = clean[2:]
                    new_page_if_needed()
                    page.insert_text((margin, y), "• ", fontsize=fontsize, fontname="helv")
                    plain = text.replace("**", "").replace("*", "")
                    for i, part in enumerate(textwrap.wrap(plain, width=chars_fit(cw - 12, fontsize))):
                        new_page_if_needed()
                        page.insert_text((margin + 12, y), part, fontsize=fontsize, fontname="helv")
                        y += line_height
                else:
                    plain = clean.replace("**", "").replace("*", "")
                    new_page_if_needed()
                    for part in textwrap.wrap(plain, width=chars_fit(cw, fontsize)):
                        new_page_if_needed()
                        page.insert_text((margin, y), part, fontsize=fontsize, fontname="helv")
                        y += line_height

            os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)
            doc.save(pdf_path)
            doc.close()
        except Exception as e:
            logger.exception("PDF generation failed for %s: %s", md_path, e)
            raise


def _add_paragraph_with_bold(doc, text: str, style=None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = p.add_run(seg)
        if i % 2 == 1:
            run.bold = True
