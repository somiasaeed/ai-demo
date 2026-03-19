"""CV Tailorer Agent: reads CV, cover letter, and job description, then produces tailored versions."""

import logging
import os
import re
import textwrap
from typing import Optional

logger = logging.getLogger(__name__)


def _add_paragraph_with_bold(doc, text: str, style: Optional[str] = None) -> None:
    """Add a paragraph to docx with **bold** segments rendered as bold runs."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = p.add_run(seg)
        if i % 2 == 1:
            run.bold = True


from agents.base import BaseAgent, AgentConfig, load_prompt
from settings import Settings
from tools.document import read_pdf_tool, read_file_tool, write_file_tool


class CVTailorerAgent(BaseAgent):
    """Agent that tailors CV and cover letter to a job description."""

    def __init__(self, settings: Optional[Settings] = None):
        config = AgentConfig(
            name="cv_tailorer",
            system_prompt=load_prompt("cv_tailorer"),
            max_tokens=4096,
        )
        super().__init__(config, settings)

    def get_tools(self) -> list:
        return [read_pdf_tool, read_file_tool, write_file_tool]

    @staticmethod
    def _next_version(output_dir: str) -> int:
        """Determine the next version number by scanning existing output files."""
        if not os.path.isdir(output_dir):
            return 1
        pattern = re.compile(r"tailored_cv_v(\d+)")
        max_ver = 0
        for name in os.listdir(output_dir):
            m = pattern.search(name)
            if m:
                max_ver = max(max_ver, int(m.group(1)))
        return max_ver + 1

    def tailor(
        self,
        cv_path: str,
        cover_letter_path: str,
        job_desc_path: str,
        output_dir: str = "output",
        photo_path: Optional[str] = None,
    ) -> str:
        """Tailor CV and cover letter to a job description.

        Auto-versions output files (v1, v2, ...) so previous runs are preserved.

        Args:
            cv_path: Path to CV PDF
            cover_letter_path: Path to cover letter PDF
            job_desc_path: Path to job description text file
            output_dir: Directory for output files
            photo_path: Optional path to photo/image to place at top-right of CV PDF

        Returns:
            Agent's summary of changes made
        """
        version = self._next_version(output_dir)
        cv_out = f"{output_dir}/tailored_cv_v{version}.md"
        cl_out = f"{output_dir}/tailored_cover_letter_v{version}.md"

        prompt = f"""\
Please tailor my CV and cover letter for the job description.

Files:
- CV: {cv_path}
- Cover Letter: {cover_letter_path}
- Job Description: {job_desc_path}

Save the tailored versions to:
- {cv_out}
- {cl_out}

Read all three files first, then produce the tailored versions.
After writing both files, provide a brief summary of the key changes you made."""

        result = self.run(prompt)

        # Convert the markdown outputs to PDF and Word
        self._md_to_pdf(cv_out, cv_out.replace(".md", ".pdf"), photo_path=photo_path)
        self._md_to_pdf(cl_out, cl_out.replace(".md", ".pdf"))
        self._md_to_docx(cv_out, cv_out.replace(".md", ".docx"), photo_path=photo_path)
        self._md_to_docx(cl_out, cl_out.replace(".md", ".docx"))

        return result

    @staticmethod
    def _normalize_date_line(line: str) -> str:
        """Fix concatenated dates to 'Month YYYY – Month YYYY' with en-dash."""
        months = (
            "January|February|March|April|May|June|July|"
            "August|September|October|November|December"
        )
        return re.sub(rf"(\d{{4}})({months})", r"\1 – \2", line)

    @staticmethod
    def _strip_link_labels(line: str) -> str:
        """Remove [LinkedIn], [github] and similar labels; keep only the link URL text."""
        line = re.sub(r"\s*\[LinkedIn\]\s*", " ", line, flags=re.IGNORECASE)
        line = re.sub(r"\s*\[github\]\s*", " ", line, flags=re.IGNORECASE)
        return re.sub(r"\s+", " ", line).strip()

    @staticmethod
    def _md_to_docx(md_path: str, docx_path: str, photo_path: Optional[str] = None) -> None:
        """Convert a markdown file to a Word document (.docx). Optionally add photo for CV."""
        if not os.path.exists(md_path):
            return
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
                try:
                    with open(md_path, encoding=enc) as f:
                        raw_lines = f.readlines()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(md_path, encoding="utf-8", errors="replace") as f:
                    raw_lines = f.readlines()
            lines = [
                CVTailorerAgent._strip_link_labels(CVTailorerAgent._normalize_date_line(ln))
                for ln in raw_lines
            ]

            doc = Document()
            style = doc.styles["Normal"]
            style.paragraph_format.space_after = Pt(6)

            if photo_path and os.path.exists(photo_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(photo_path, width=Pt(90), height=Pt(110))
                except Exception:
                    pass

            for line in lines:
                clean = line.rstrip("\n")
                if clean.startswith("# "):
                    doc.add_heading(clean[2:].strip(), level=0)
                elif clean.startswith("## "):
                    doc.add_heading(clean[3:].strip().upper().replace("**", ""), level=1)
                elif clean.startswith("### "):
                    rest = clean[4:].replace("**", "").strip()
                    if "\t" in rest:
                        title_part, location_part = rest.split("\t", 1)
                        p = doc.add_paragraph()
                        run = p.add_run(title_part.strip() + "  ")
                        run.bold = True
                        run = p.add_run(location_part.strip())
                        run.bold = True
                        p.paragraph_format.space_before = Pt(10)
                    else:
                        p = doc.add_paragraph(rest)
                        p.paragraph_format.space_before = Pt(10)
                        for run in p.runs:
                            run.bold = True
                elif clean.startswith("---"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(8)
                elif clean.startswith("- ") or clean.startswith("* "):
                    text = clean[2:]
                    _add_paragraph_with_bold(doc, text, style="List Bullet")
                else:
                    plain = clean.replace("**", "").replace("*", "").strip()
                    if plain.startswith("Tech Stack:"):
                        p = doc.add_paragraph(plain)
                        for run in p.runs:
                            run.italic = True
                        p.paragraph_format.space_after = Pt(20)
                    else:
                        _add_paragraph_with_bold(doc, clean, style=None)

            out_dir = os.path.dirname(docx_path)
            if out_dir:
                os.makedirs(out_dir, exist_ok=True)
            doc.save(docx_path)
        except Exception as e:
            logger.exception("Word generation failed for %s: %s", md_path, e)
            raise

    @staticmethod
    def _md_to_pdf(md_path: str, pdf_path: str, photo_path: Optional[str] = None) -> None:
        """Convert a markdown file to a simple PDF. If photo_path is set (for CV), place image at top right of first page."""
        if not os.path.exists(md_path):
            return
        try:
            import pymupdf

            with open(md_path) as f:
                raw_lines = f.readlines()
            lines = [
                CVTailorerAgent._strip_link_labels(CVTailorerAgent._normalize_date_line(ln))
                for ln in raw_lines
            ]

            doc = pymupdf.open()
            fontsize = 9.5
            line_height = fontsize * 1.4
            margin = 50
            page_w, page_h = 595, 842  # A4
            bottom_y = page_h - margin

            # Photo box: top-right corner, aligned with header (name/title/contact)
            img_w, img_h = 90, 110
            photo_top = margin - 8
            photo_rect = pymupdf.Rect(page_w - margin - img_w, photo_top, page_w - margin, photo_top + img_h)
            # First separator line must sit below the whole header block (text + image)
            first_separator_min_y = photo_top + img_h + 12
            content_width_no_photo = page_w - 2 * margin
            content_width_with_photo = page_w - margin - img_w - 15 - margin

            def chars_that_fit(width_pt: float, size: float) -> int:
                """Approximate max characters per line for given width (Helvetica)."""
                avg = 0.55 * size
                return max(20, int(width_pt / avg))

            def draw_text_with_bold(
                get_page,
                x_left: float,
                x_right: float,
                y_start: float,
                line_height: float,
                fontsize: float,
                text: str,
                indent_pt: float = 0,
            ) -> float:
                """Draw text that may contain **bold** segments; word-wrap and return final y."""
                parts = text.split("**")
                segments = [(parts[i], (i % 2) == 1) for i in range(len(parts))]
                x, current_y = x_left + indent_pt, y_start
                for seg_text, is_bold in segments:
                    font = "hebo" if is_bold else "helv"  # hebo = Helvetica-Bold (PyMuPDF Base14)
                    for word in seg_text.split():
                        w = word + " "
                        w_len = pymupdf.get_text_length(w, fontname=font, fontsize=fontsize)
                        if x + w_len > x_right and x > x_left + indent_pt:
                            current_y += line_height
                            x = x_left + indent_pt
                            if current_y >= bottom_y:
                                maybe_new_page()
                                current_y = margin
                                x = x_left + indent_pt
                        page = get_page()
                        page.insert_text((x, current_y), w, fontsize=fontsize, fontname=font)
                        x += w_len
                return current_y + line_height

            # Consistent bullet: symbol + indent (pt) so all list items align
            bullet_indent_pt = 12
            bullet_char = "• "
            job_entry_spacing_pt = 14  # space between job entries (padding so text doesn't bleed)
            section_margin_bottom_pt = 14  # margin below each major section (Summary, Experience, Skills)
            section_heading_margin_bottom_pt = 8  # space below ## heading before content
            padding_after_job_title_pt = 6  # space below job title line before date/bullets
            padding_after_rule_pt = 6  # space below --- before next block
            project_block_gap_pt = 20  # gap between project blocks (Projects section)

            page = None
            page_num = 0
            first_heading3 = True
            first_section = True

            y = margin
            for line in lines:
                clean = line.rstrip("\n")
                if page is None:
                    page = doc.new_page(width=page_w, height=page_h)
                    page_num = 1
                    y = margin
                    if photo_path and os.path.exists(photo_path):
                        page.insert_image(photo_rect, filename=photo_path, keep_proportion=True, overlay=False)
                has_photo = photo_path and os.path.exists(photo_path)
                is_first_page = page_num == 1
                content_width = content_width_with_photo if (is_first_page and has_photo) else content_width_no_photo
                right = (margin + content_width) if (is_first_page and has_photo) else (page_w - margin)

                def maybe_new_page() -> None:
                    nonlocal page, page_num, y, content_width, right
                    if y >= bottom_y:
                        page = doc.new_page(width=page_w, height=page_h)
                        page_num += 1
                        y = margin
                        content_width = content_width_no_photo
                        right = page_w - margin

                if clean.startswith("# "):
                    text = clean[2:]
                    width_chars = chars_that_fit(content_width, 14)
                    for part in textwrap.wrap(text, width=width_chars):
                        maybe_new_page()
                        page.insert_text((margin, y), part, fontsize=14, fontname="helv")
                        y += line_height
                elif clean.startswith("## "):
                    # Section headings: ALL CAPS, bold; margin above (except first) so sections stack clearly
                    maybe_new_page()
                    if not first_section:
                        y += section_margin_bottom_pt
                    first_section = False
                    text = clean[3:].strip().upper().replace("**", "")
                    width_chars = chars_that_fit(content_width, 13)
                    for part in textwrap.wrap(text, width=width_chars):
                        maybe_new_page()
                        page.insert_text((margin, y), part, fontsize=13, fontname="hebo")
                        y += line_height
                    y += section_heading_margin_bottom_pt
                elif clean.startswith("### "):
                    # Job title line: bold; optional tab-separated location right-aligned
                    maybe_new_page()
                    if not first_heading3:
                        y += job_entry_spacing_pt
                    first_heading3 = False
                    rest = clean[4:]
                    if "\t" in rest:
                        title_part, location_part = rest.split("\t", 1)
                        title_part = title_part.replace("**", "").strip()
                        location_part = location_part.replace("**", "").strip()
                        page.insert_text((margin, y), title_part, fontsize=10.5, fontname="hebo")
                        loc_w = pymupdf.get_text_length(location_part, fontname="hebo", fontsize=10.5)
                        page.insert_text((right - loc_w, y), location_part, fontsize=10.5, fontname="hebo")
                        y += line_height + padding_after_job_title_pt
                    else:
                        text = rest.replace("**", "").strip()
                        width_chars = chars_that_fit(content_width, 10.5)
                        for part in textwrap.wrap(text, width=width_chars):
                            maybe_new_page()
                            page.insert_text((margin, y), part, fontsize=10.5, fontname="hebo")
                            y += line_height
                        y += padding_after_job_title_pt
                elif clean.startswith("---"):
                    maybe_new_page()
                    # Keep first separator below header block (name + contact + image) when photo is present
                    if is_first_page and has_photo:
                        y = max(y, first_separator_min_y)
                    page.draw_line((margin, y - 2), (right, y - 2))
                    y += line_height + padding_after_rule_pt
                else:
                    # Bullet lines: normalize "- " or "* " to consistent • and indent
                    is_bullet = clean.startswith("- ") or clean.startswith("* ")
                    if is_bullet:
                        text = clean[2:]
                        maybe_new_page()
                        if "**" in text:
                            page.insert_text((margin, y), bullet_char, fontsize=fontsize, fontname="helv")
                            y = draw_text_with_bold(
                                lambda: page, margin, right, y, line_height, fontsize, text, bullet_indent_pt
                            )
                        else:
                            text_plain = text.replace("**", "").replace("*", "")
                            width_for_wrap = content_width - bullet_indent_pt
                            width_chars = chars_that_fit(width_for_wrap, fontsize)
                            parts = textwrap.wrap(text_plain, width=width_chars)
                            for i, part in enumerate(parts):
                                maybe_new_page()
                                if i == 0:
                                    page.insert_text((margin, y), bullet_char, fontsize=fontsize, fontname="helv")
                                    page.insert_text((margin + bullet_indent_pt, y), part, fontsize=fontsize, fontname="helv")
                                else:
                                    page.insert_text((margin + bullet_indent_pt, y), part, fontsize=fontsize, fontname="helv")
                                y += line_height
                    else:
                        text = clean
                        maybe_new_page()
                        text_plain = text.replace("**", "").replace("*", "")
                        is_tech_stack = text_plain.strip().startswith("Tech Stack:")
                        if "**" in text:
                            y = draw_text_with_bold(
                                lambda: page, margin, right, y, line_height, fontsize, text, 0
                            )
                        elif is_tech_stack:
                            # Projects section: Tech Stack line in italics (heit)
                            width_chars = chars_that_fit(content_width, fontsize)
                            for part in textwrap.wrap(text_plain.strip(), width=width_chars):
                                maybe_new_page()
                                page.insert_text((margin, y), part, fontsize=fontsize, fontname="heit")
                                y += line_height
                            y += project_block_gap_pt
                        else:
                            width_chars = chars_that_fit(content_width, fontsize)
                            for part in textwrap.wrap(text_plain, width=width_chars):
                                maybe_new_page()
                                page.insert_text((margin, y), part, fontsize=fontsize, fontname="helv")
                                y += line_height

            out_dir = os.path.dirname(pdf_path)
            if out_dir:
                os.makedirs(out_dir, exist_ok=True)
            doc.save(pdf_path)
            doc.close()
        except Exception as e:
            logger.exception("PDF generation failed for %s: %s", md_path, e)
            raise
